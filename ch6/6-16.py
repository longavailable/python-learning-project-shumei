# -*- coding: utf-8 -*-
"""
Created on Sun Aug 24 12:18:37 2025

@author: Xiaolong Liu
"""

from docx import Document

doc = Document()

p = doc.add_paragraph('细雨潇潇欲晓天。')

p = doc.add_paragraph()

p.add_run('---')

p.add_run('张坚').bold = True

p.add_run('《偶成》').italic = True

doc.save('6-16.docx')
